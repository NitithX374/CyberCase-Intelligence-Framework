from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "CyberCase_Thesis_V1.docx"
FONT = "Angsana New"
MONO = "Consolas"
INK = RGBColor(31, 45, 61)
ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(92, 103, 112)
AMBER = RGBColor(151, 92, 0)
CHAPTERS = [ROOT / f"0{i}_{name}.md" for i, name in enumerate([
    "chapter_1_introduction", "chapter_2_related_theory", "chapter_3_methodology",
    "chapter_4_implementation", "chapter_5_results", "chapter_6_conclusion"
], 1)]


def set_font(run, name=FONT, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade(element, fill):
    properties = element.get_or_add_pPr() if element.tag.endswith("}p") else element.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            set_cell_margins(cell)


def add_field(paragraph, instruction, display=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(16)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    specs = {
        "Title": (28, ACCENT, 0, 12), "Subtitle": (18, MUTED, 0, 8),
        "Heading 1": (22, ACCENT, 18, 10), "Heading 2": (19, ACCENT, 14, 7),
        "Heading 3": (17, INK, 10, 5), "Caption": (14, MUTED, 8, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(16)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15
    code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = MONO
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Cm(0.45)
    code.paragraph_format.right_indent = Cm(0.45)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    callout.font.name = FONT
    callout.font.size = Pt(15)
    callout.font.color.rgb = AMBER
    callout.paragraph_format.left_indent = Cm(0.5)
    callout.paragraph_format.right_indent = Cm(0.5)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(7)


def add_inline(paragraph, text, citation_numbers):
    token_pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\[@.+?\]|\[[^\]]+\]\([^)]+\))")
    cursor = 0
    for match in token_pattern.finditer(text):
        if match.start() > cursor:
            set_font(paragraph.add_run(text[cursor:match.start()]), size=16)
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=16, bold=True)
        elif token.startswith("`"):
            set_font(paragraph.add_run(token[1:-1]), name=MONO, size=10, color=ACCENT)
        elif token.startswith("[@"):
            keys = re.findall(r"@([A-Za-z0-9_-]+)", token)
            labels = [str(citation_numbers[key]) for key in keys if key in citation_numbers]
            set_font(paragraph.add_run("[" + ", ".join(labels) + "]"), size=14, color=ACCENT)
        else:
            label, target = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            set_font(paragraph.add_run(f"{label} ({target})"), size=15, color=ACCENT)
        cursor = match.end()
    if cursor < len(text):
        set_font(paragraph.add_run(text[cursor:]), size=16)
    if re.search(r"\[(TODO|VERIFY|RESULT PENDING|FIGURE NEEDED)", text):
        shade(paragraph._p, "FFF4CE")


def table_widths(rows, total=9360):
    columns = len(rows[0])
    scores = [max(8, max(len(row[i]) for row in rows)) for i in range(columns)]
    capped = [min(score, 58) for score in scores]
    widths = [max(1050, round(total * score / sum(capped))) for score in capped]
    widths[-1] += total - sum(widths)
    return widths


def add_table(doc, rows, citation_numbers):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, table_widths(rows))
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 or len(value) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(paragraph, value, citation_numbers)
            for run in paragraph.runs:
                set_font(run, size=13, bold=row_index == 0)
            if row_index == 0:
                shade(cell._tc, "E8EEF5")
                cell._tc.get_or_add_tcPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_markdown(doc, text, citation_numbers, skip_h1=False):
    lines = text.splitlines()
    index = 0
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index].rstrip()
        if line.startswith("```"):
            if in_code:
                paragraph = doc.add_paragraph(style="Code Block")
                shade(paragraph._p, "F2F4F7")
                run = paragraph.add_run("\n".join(code_lines))
                set_font(run, name=MONO, size=9, color=INK)
                code_lines = []
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[index + 1]):
            rows = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                parts = [part.strip() for part in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", part.replace(" ", "")) for part in parts):
                    rows.append(parts)
                index += 1
            add_table(doc, rows, citation_numbers)
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            level, title = len(heading.group(1)), heading.group(2)
            if level == 1 and skip_h1:
                index += 1
                continue
            if title.startswith("ตารางที่"):
                paragraph = doc.add_paragraph(style="Caption")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            add_inline(paragraph, title, citation_numbers)
            index += 1
            continue
        if line.startswith("> "):
            paragraph = doc.add_paragraph(style="Callout")
            shade(paragraph._p, "FFF4CE")
            add_inline(paragraph, line[2:], citation_numbers)
            index += 1
            continue
        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if bullet or numbered:
            paragraph = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            add_inline(paragraph, (bullet or numbered).group(1), citation_numbers)
            index += 1
            continue
        if line == "---":
            index += 1
            continue
        if line:
            paragraph = doc.add_paragraph()
            add_inline(paragraph, line.replace("  ", " "), citation_numbers)
        index += 1


def bib_entries(text):
    entries = {}
    for block in re.split(r"\n(?=@)", text.strip()):
        first = re.match(r"@\w+\{([^,]+),", block)
        if not first:
            continue
        fields = {}
        for match in re.finditer(r"(?m)^\s*(\w+)\s*=\s*\{(.*?)\}\s*,?\s*$", block):
            fields[match.group(1).lower()] = match.group(2).replace(r"\&", "&")
        entries[first.group(1)] = fields
    return entries


def format_reference(fields):
    authors = fields.get("author", "Unknown author").replace("{{", "").replace("}}", "")
    title = fields.get("title", "Untitled")
    year = fields.get("year", "n.d.")
    venue = fields.get("journal") or fields.get("booktitle") or fields.get("institution") or ""
    details = ", ".join(part for part in [venue, fields.get("volume"), fields.get("pages")] if part)
    locator = fields.get("doi") or fields.get("url", "")
    return f"{authors}. ({year}). {title}. {details}. {locator}".replace("..", ".")


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("วิทยานิพนธ์ฉบับร่าง • VERSION 1"), size=15, bold=True, color=AMBER)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(28)
    title.paragraph_format.space_after = Pt(14)
    set_font(title.add_run("กรอบงานไซเบอร์เคสสำหรับช่วยวิเคราะห์ข้อมูลคดี\nด้วยปัญญาประดิษฐ์โดยคงการอ้างอิงแหล่งข้อมูล"), size=26, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.add_run("CyberCase Intelligence Framework:\nAn Evidence-Grounded AI-Assisted Case Analysis System"), size=17, color=ACCENT)
    for _ in range(4):
        doc.add_paragraph()
    for value in ("จัดทำโดย [ชื่อนักศึกษา — PLACEHOLDER]", "รหัสนักศึกษา [PLACEHOLDER]", "อาจารย์ที่ปรึกษา [PLACEHOLDER]", "มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ [VERIFY]", "ปีการศึกษา [PLACEHOLDER]"):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        set_font(paragraph.add_run(value), size=16, color=INK)
        if "[" in value:
            shade(paragraph._p, "FFF4CE")
    doc.add_page_break()


def add_toc(doc):
    heading = doc.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(heading.add_run("สารบัญ"), size=22, bold=True, color=ACCENT)
    paragraph = doc.add_paragraph()
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u', "สารบัญ")
    document_settings = doc.settings._element
    update = document_settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        document_settings.append(update)
    update.set(qn("w:val"), "true")
    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True
    configure_styles(doc)
    doc.core_properties.title = "CyberCase Intelligence Framework Thesis Version 1"
    doc.core_properties.subject = "Evidence-grounded AI-assisted case analysis"
    doc.core_properties.author = "CyberCase project team"
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("CYBERCASE INTELLIGENCE FRAMEWORK • THESIS V1"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("CyberCase Thesis V1  •  "), size=9, color=MUTED)
    add_field(footer, "PAGE", "1")
    all_text = "\n".join(path.read_text(encoding="utf-8") for path in CHAPTERS)
    citation_keys = []
    for key in re.findall(r"@([A-Za-z0-9_-]+)", all_text):
        if key not in citation_keys:
            citation_keys.append(key)
    citation_numbers = {key: index + 1 for index, key in enumerate(citation_keys)}
    add_cover(doc)
    front = (ROOT / "00_frontmatter.md").read_text(encoding="utf-8")
    sections = {name.strip(): body for name, body in re.findall(r"(?ms)^##\s+([^\n]+)\n(.*?)(?=^##\s+|\Z)", front)}
    for source_name, display_name in (("บทคัดย่อ", "บทคัดย่อ"), ("Abstract", "Abstract"), ("กิตติกรรมประกาศ", "กิตติกรรมประกาศ")):
        heading = doc.add_paragraph(style="Heading 1")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(heading, display_name, citation_numbers)
        add_markdown(doc, sections[source_name].strip(), citation_numbers)
        doc.add_page_break()
    add_toc(doc)
    for source_name, display_name in (("สารบัญภาพฉบับร่าง", "สารบัญภาพ"), ("สารบัญตารางฉบับร่าง", "สารบัญตาราง")):
        heading = doc.add_paragraph(style="Heading 1")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(heading, display_name, citation_numbers)
        add_markdown(doc, sections[source_name].strip(), citation_numbers)
        doc.add_page_break()
    for chapter in CHAPTERS:
        text = chapter.read_text(encoding="utf-8")
        title = re.match(r"^#\s+(.+)$", text, re.MULTILINE).group(1)
        heading = doc.add_paragraph(style="Heading 1")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(heading, title, citation_numbers)
        add_markdown(doc, text, citation_numbers, skip_h1=True)
        doc.add_page_break()
    bibliography = bib_entries((ROOT / "references.bib").read_text(encoding="utf-8"))
    heading = doc.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(heading, "บรรณานุกรม", citation_numbers)
    for key in citation_keys:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Cm(1.0)
        paragraph.paragraph_format.first_line_indent = Cm(-0.6)
        add_inline(paragraph, format_reference(bibliography.get(key, {})), citation_numbers)
    while doc.paragraphs and not doc.paragraphs[-1].text:
        element = doc.paragraphs[-1]._element
        element.getparent().remove(element)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
